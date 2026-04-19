from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

_SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "summary": (
        "自我评价",
        "自我介绍",
        "个人简介",
        "个人总结",
        "summary",
        "profile",
    ),
    "education": ("教育经历", "教育背景", "学历", "education"),
    "experience": (
        "工作经历",
        "工作经验",
        "实习经历",
        "experience",
        "employment",
    ),
    "projects": ("项目经历", "项目经验", "project", "projects"),
    "awards": ("获奖经历", "奖项", "荣誉", "awards", "honors"),
    "target_position": (
        "目标职位",
        "目标岗位",
        "应聘职位",
        "应聘岗位",
        "求职意向",
        "target position",
        "desired position",
    ),
    "skills": ("专业技能", "技能", "技术栈", "skill", "skills"),
}

_FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "name": ("姓名", "name"),
    "gender": ("性别", "gender"),
    "major": ("专业", "major"),
    "education_level": ("学历", "最高学历", "degree", "education level"),
    "self_introduction": (
        "自我介绍",
        "个人简介",
        "个人总结",
        "summary",
        "profile",
    ),
    "target_position": (
        "目标职位",
        "目标岗位",
        "应聘职位",
        "应聘岗位",
        "求职意向",
        "target position",
        "desired position",
    ),
}

_MAX_SECTION_ITEMS = 6
_MAX_SUMMARY_LENGTH = 500
_DEFAULT_SUMMARY = "简历已上传，暂未解析。"


def read_resume_text(path: str) -> str:
    file_path = Path(path)
    if not file_path.exists():
        return ""

    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                return "\n".join((page.extract_text() or "") for page in pdf.pages)
        except Exception:
            return ""

    if suffix in {".docx", ".doc"}:
        return _read_docx_or_doc(file_path)

    return ""


def _read_docx_or_doc(file_path: Path) -> str:
    """Read .docx and .doc files and extract text."""
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        try:
            import docx

            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)
        except Exception as exc:
            logger.warning("Failed to read .docx file %s: %s", file_path, exc)
            return ""

    if suffix == ".doc":
        # Try python-docx first (for some .doc files that are actually DOCX)
        try:
            import docx

            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            if paragraphs:
                return "\n".join(paragraphs)
        except Exception:
            pass

        # Try textract as fallback
        try:
            import textract

            text = textract.process(str(file_path), method="antiword")
            return text.decode("utf-8", errors="ignore").strip()
        except Exception:
            pass

        # Try textract with other methods
        for method in ["catdoc", "wvText"]:
            try:
                import subprocess

                result = subprocess.run(
                    [method, str(file_path)],
                    capture_output=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    return result.stdout.decode("utf-8", errors="ignore").strip()
            except Exception:
                continue

        logger.warning("No suitable converter found for .doc file %s", file_path)
        return ""

    return ""


def parse_resume_text(text: str) -> dict[str, Any]:
    cleaned_text = _clean_text(text)
    if not cleaned_text:
        return _empty_resume(_DEFAULT_SUMMARY)

    scalar_fields: dict[str, str] = {
        "name": "",
        "gender": "",
        "major": "",
        "education_level": "",
        "self_introduction": "",
        "target_position": "",
    }

    sections: dict[str, list[str]] = {
        "education": [],
        "experience": [],
        "projects": [],
        "awards": [],
        "skills": [],
    }
    summary_lines: list[str] = []
    lead_lines: list[str] = []
    current_section: str | None = None

    for raw_line in cleaned_text.splitlines():
        line = _clean_line(raw_line)
        if not line:
            continue

        scalar = _extract_scalar_field(line)
        if scalar:
            key, value = scalar
            if key in scalar_fields and value and not scalar_fields[key]:
                scalar_fields[key] = value
            if key == "self_introduction" and value:
                _append_unique(summary_lines, value)
            continue

        detected = _detect_section(line)
        if detected:
            if detected == "target_position":
                current_section = None
                inline_target = _extract_inline_content(line)
                if inline_target and not scalar_fields["target_position"]:
                    scalar_fields["target_position"] = inline_target
                continue

            current_section = (
                detected if detected in sections or detected == "summary" else None
            )
            inline = _extract_inline_content(line)
            if inline:
                if current_section == "summary":
                    _append_unique(summary_lines, inline)
                    if not scalar_fields["self_introduction"]:
                        scalar_fields["self_introduction"] = inline
                elif current_section in sections:
                    _append_unique(sections[current_section], inline)
            continue

        if current_section == "summary":
            _append_unique(summary_lines, line)
            continue

        if current_section in sections:
            _append_unique(sections[current_section], line)
            continue

        _append_unique(lead_lines, line)

    summary = _build_summary(cleaned_text, summary_lines, lead_lines)

    if not scalar_fields["self_introduction"]:
        scalar_fields["self_introduction"] = summary

    if not scalar_fields["education_level"]:
        scalar_fields["education_level"] = _infer_education_level(sections["education"])

    if not scalar_fields["target_position"]:
        scalar_fields["target_position"] = _infer_target_position(
            skills=sections["skills"],
            projects=sections["projects"],
            experience=sections["experience"],
        )

    return {
        "name": scalar_fields["name"],
        "gender": scalar_fields["gender"],
        "major": scalar_fields["major"],
        "education_level": scalar_fields["education_level"],
        "self_introduction": scalar_fields["self_introduction"],
        "summary": summary,
        "education": sections["education"][:_MAX_SECTION_ITEMS],
        "experience": sections["experience"][:_MAX_SECTION_ITEMS],
        "projects": sections["projects"][:_MAX_SECTION_ITEMS],
        "awards": sections["awards"][:_MAX_SECTION_ITEMS],
        "target_position": scalar_fields["target_position"],
        "skills": sections["skills"][:_MAX_SECTION_ITEMS],
    }


def parse_resume(path: str) -> dict[str, Any]:
    file_path = Path(path)
    if not file_path.exists():
        return _empty_resume("简历文件不存在")

    text = read_resume_text(path)
    if not text.strip():
        return _empty_resume(_DEFAULT_SUMMARY)

    return parse_resume_text(text)


def build_resume_prompt(parsed: dict[str, Any] | None) -> str:
    if not parsed:
        return "暂无简历"

    name = str(parsed.get("name") or "").strip()
    major = str(parsed.get("major") or "").strip()
    education_level = str(parsed.get("education_level") or "").strip()
    target_position = str(parsed.get("target_position") or "").strip()
    self_intro = str(parsed.get("self_introduction") or "").strip()
    summary = str(parsed.get("summary") or "").strip()
    projects = _normalize_items(parsed.get("projects"))
    awards = _normalize_items(parsed.get("awards"))
    skills = _normalize_items(parsed.get("skills"))

    parts: list[str] = []
    if name:
        parts.append(f"候选人姓名: {name}")
    if major or education_level or target_position:
        tags = [
            f"专业: {major}" if major else "",
            f"学历: {education_level}" if education_level else "",
            f"目标职位: {target_position}" if target_position else "",
        ]
        compact = "；".join(item for item in tags if item)
        if compact:
            parts.append(compact)
    if self_intro and self_intro != summary:
        parts.append("自我介绍: " + self_intro)
    if summary:
        parts.append(summary)
    if projects:
        parts.append("项目经历: " + "；".join(projects[:3]))
    if awards:
        parts.append("获奖经历: " + "；".join(awards[:3]))
    if skills:
        parts.append("核心技能: " + "、".join(skills[:6]))

    return "\n".join(parts) or "暂无简历"


def _clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _clean_line(value: str) -> str:
    cleaned = re.sub(r"\s+", " ", value).strip()
    cleaned = cleaned.strip("-•·\t ")
    return cleaned


def _normalize_heading(value: str) -> str:
    compact = value.lower()
    compact = re.sub(r"[\s：:()（）\[\]【】|/\\-]+", "", compact)
    return compact


def _detect_section(line: str) -> str | None:
    heading = _normalize_heading(line)
    if not heading or len(heading) > 30:
        return None

    for section, aliases in _SECTION_ALIASES.items():
        for alias in aliases:
            if _normalize_heading(alias) in heading:
                return section
    return None


def _extract_inline_content(line: str) -> str:
    if ":" in line:
        return line.split(":", 1)[1].strip()
    if "：" in line:
        return line.split("：", 1)[1].strip()
    return ""


def _append_unique(bucket: list[str], value: str) -> None:
    cleaned = _clean_line(value)
    if not cleaned:
        return
    if len(cleaned) > 180:
        cleaned = cleaned[:180]
    if cleaned not in bucket:
        bucket.append(cleaned)


def _build_summary(
    full_text: str,
    summary_lines: list[str],
    lead_lines: list[str],
) -> str:
    if summary_lines:
        summary = "；".join(summary_lines[:3])
    elif lead_lines:
        summary = "；".join(lead_lines[:3])
    else:
        summary = full_text.replace("\n", " ")

    summary = re.sub(r"\s+", " ", summary).strip()
    if not summary:
        return _DEFAULT_SUMMARY
    return summary[:_MAX_SUMMARY_LENGTH]


def _extract_scalar_field(line: str) -> tuple[str, str] | None:
    if not line or (":" not in line and "：" not in line):
        return None

    if ":" in line:
        key_raw, value_raw = line.split(":", 1)
    else:
        key_raw, value_raw = line.split("：", 1)

    key = _normalize_heading(key_raw)
    value = _clean_line(value_raw)
    if not key or not value:
        return None

    for field_name, aliases in _FIELD_ALIASES.items():
        for alias in aliases:
            if _normalize_heading(alias) == key:
                return field_name, value

    return None


def _infer_education_level(education_items: list[str]) -> str:
    text = " ".join(education_items)
    checks = (
        ("博士", "博士"),
        ("硕士", "硕士"),
        ("研究生", "硕士/研究生"),
        ("本科", "本科"),
        ("学士", "本科"),
        ("大专", "大专"),
        ("专科", "大专"),
        ("中专", "中专"),
    )
    for token, label in checks:
        if token in text:
            return label
    return ""


def _infer_target_position(
    *,
    skills: list[str],
    projects: list[str],
    experience: list[str],
) -> str:
    corpus = " ".join(skills + projects + experience).lower()
    mapping = (
        ("前端", "前端工程师"),
        ("frontend", "前端工程师"),
        ("react", "前端工程师"),
        ("后端", "后端工程师"),
        ("backend", "后端工程师"),
        ("java", "后端工程师"),
        ("python", "算法/后端工程师"),
        ("算法", "算法工程师"),
        ("机器学习", "算法工程师"),
        ("测试", "测试工程师"),
        ("qa", "测试工程师"),
        ("运维", "运维/SRE"),
        ("sre", "运维/SRE"),
        ("产品", "产品经理"),
        ("数据分析", "数据分析师"),
    )
    for token, role in mapping:
        if token in corpus:
            return role
    return ""


def _normalize_items(value: Any) -> list[str]:
    if isinstance(value, list):
        items = value
    elif isinstance(value, str):
        items = [value]
    else:
        return []

    output: list[str] = []
    for item in items:
        text = _clean_line(str(item))
        if not text or text in output:
            continue
        output.append(text)
    return output


def _empty_resume(summary: str) -> dict[str, Any]:
    return {
        "name": "",
        "gender": "",
        "major": "",
        "education_level": "",
        "self_introduction": summary,
        "summary": summary,
        "education": [],
        "experience": [],
        "projects": [],
        "awards": [],
        "target_position": "",
        "skills": [],
    }
